"""Generate thesis section 7.3 as a Word document with proper formatting."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h.font.size = Pt(16)
    elif level == 2:
        h.font.size = Pt(14)
    elif level == 3:
        h.font.size = Pt(13)
    else:
        h.font.size = Pt(12)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_formula(text, label=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    if label:
        run2 = p.add_run(f"    ({label})")
        run2.font.size = Pt(11)
    return p

def add_table(headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()  # spacing after table
    return table

def add_figure_placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"[{text}]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.right_indent = Cm(1.27)
    run = p.add_run(text)
    run.italic = True
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ============================================================
# SECTION 7.3
# ============================================================

add_heading("7.3 Core Functionalities Implementation", level=1)

add_para("This section presents the design, implementation, and evaluation of SemanticEnergy, a local hallucination detection system for large language model (LLM) outputs. The discussion follows a narrative structure: beginning with the theoretical foundations, progressing through dataset construction and probe training, and culminating in the system architecture and integration that delivers real-time hallucination scoring.")

# --- 7.3.1 ---
add_heading("7.3.1 Foundational Frameworks: Semantic Entropy and Semantic Energy", level=2)

add_para("The detection of hallucinations in LLM outputs requires a principled measure of model uncertainty. Two complementary frameworks provide this foundation: Semantic Entropy and Semantic Energy.")

add_heading("Semantic Entropy", level=3)

add_para("Semantic Entropy, introduced by Kuhn et al. (2023), measures uncertainty by examining whether a model produces consistent answers across multiple independent generations. The procedure is as follows:")

add_bullet("Generate N independent responses to the same prompt (using stochastic sampling).")
add_bullet("Cluster the responses by semantic equivalence — grouping answers that convey the same meaning, regardless of exact wording.")
add_bullet("Compute the Shannon entropy over the cluster size distribution.")

add_para("If all responses converge to a single meaning, entropy is near zero (high certainty). If responses scatter across many distinct meanings, entropy is high (significant uncertainty, elevated hallucination risk).")

add_formula("H = −Σ (|c| / N) · ln(|c| / N)", "Formula 1")
add_para("where |c| is the number of samples in cluster c, and N is the total number of generated samples.", italic=True)

add_heading("Semantic Energy", level=3)

add_para("Semantic Energy extends the entropy approach by incorporating token-level logit magnitudes — a measure of how confident the model was about each individual token it generated. The analogy is drawn from statistical mechanics: in a physical system at low temperature, particles settle into a single low-energy state (ordered, certain); at high temperature, they scatter across many states (disordered, uncertain).")

add_para("Rather than merely counting how many samples fall into each cluster, Semantic Energy weights each sample by the model's per-token confidence, then aggregates these weights by cluster.")

add_formula("E_i = −(1 / K_i) · Σ ℓ_{i,t}  for t = 1 to K_i", "Formula 2")
add_para("where ℓ_{i,t} is the raw logit of the chosen token at generation step t for sample i, and K_i is the number of tokens generated. A more negative energy indicates higher token-level confidence.", italic=True)

add_formula("E_cluster(c) = −Σ E_i  for i ∈ c", "Formula 3a")
add_formula("energy_score(c) = E_cluster(c) / Σ E_cluster(c')", "Formula 3b")
add_para("The final confidence score for the main answer is the normalised energy share of its cluster — a value in [0, 1] where higher indicates greater confidence.", italic=True)

add_heading("Comparison of the Two Frameworks", level=3)

add_table(
    ["Property", "Semantic Entropy", "Semantic Energy"],
    [
        ["Input signal", "Cluster sizes only", "Cluster sizes and per-token logits"],
        ["Output range", "[0, ln N]", "[0, 1]"],
        ["Orientation", "Higher = more uncertain", "Higher = more confident"],
        ["Sensitive to token quality", "No", "Yes — high-logit tokens increase cluster weight"],
        ["Correlation with correctness", "Negative", "Positive"],
    ],
    "Table 7.1 — Comparison of Semantic Entropy and Semantic Energy frameworks."
)

add_para("Both frameworks share the same clustering step but produce complementary views of model certainty: entropy captures structural uncertainty (how many distinct meanings exist), while energy captures both structural and qualitative uncertainty (how confidently each meaning was produced).")

add_figure_placeholder("Figure 7.1 — Process diagram: Five independent generations for a single prompt are passed through pairwise semantic equivalence checking and grouped into clusters. Two parallel teacher signals are computed from the same clustering: the energy teacher aggregates per-token logit magnitudes by cluster, while the entropy teacher computes Shannon entropy over cluster sizes.")

# --- 7.3.2 ---
add_heading("7.3.2 The Computational Cost Problem and Motivation for Linear Probes", level=2)

add_para("The full Semantic Energy pipeline, while theoretically sound, is computationally prohibitive for real-time applications.")

add_table(
    ["Operation", "Forward Passes", "Typical Latency (RTX 3060)"],
    [
        ["Generate N = 5 responses", "5", "~40s"],
        ["Pairwise semantic equivalence checks", "Up to N(N−1)/2 = 10", "~30s"],
        ["Energy and entropy computation", "Arithmetic only", "<1ms"],
        ["Total", "Up to 15", "60–120s"],
    ],
    "Table 7.2 — Computational cost breakdown of the full Semantic Energy pipeline."
)

add_para("A 60–120 second latency per query is unacceptable for interactive use. This motivates the central research question of the probe-based approach:")

add_blockquote("Can a single forward pass through the model's internal hidden states predict what the full multi-sample pipeline would have computed?")

add_para("The hypothesis rests on a key insight: when an LLM processes a question, its internal representations (hidden states) already encode whether it possesses strong knowledge about the answer. The full pipeline makes this confidence visible through expensive behavioural analysis; a linear probe can learn to read the same signal directly from hidden states, bypassing the multi-sample generation and clustering entirely.")

add_para("Two token positions yield two probe families:")
add_bullet("Token Before Generation (TBG): ", bold_prefix="")
add_para("The hidden state at the last prompt token, captured before any answer is generated. This enables pre-generation risk screening.", indent=True)
add_bullet("Second-to-Last Token (SLT): ", bold_prefix="")
add_para("The hidden state at the second-to-last generated token, captured after an answer has been produced. This enables post-generation confidence scoring with per-sentence granularity.", indent=True)

add_para("Two teacher signals (energy and entropy) combined with two token positions produce a matrix of four probes.")

add_table(
    ["Probe", "Token Position", "Teacher Signal", "Availability", "Primary Use Case"],
    [
        ["TBG Energy", "Last prompt token", "Binarised energy score", "Before generation", "Ultra-fast pre-screening"],
        ["TBG Entropy", "Last prompt token", "Binarised entropy score", "Before generation", "Ultra-fast pre-screening"],
        ["SLT Energy", "2nd-to-last generated token", "Binarised energy score", "After generation", "Per-sentence scoring"],
        ["SLT Entropy", "2nd-to-last generated token", "Binarised entropy score", "After generation", "Per-sentence scoring"],
    ],
    "Table 7.3 — The four linear probes and their characteristics."
)

# --- 7.3.3 ---
add_heading("7.3.3 Dataset Construction", level=2)

add_heading("7.3.3.1 TriviaQA as Source Corpus", level=3)

add_para("The training dataset is derived from TriviaQA (Joshi et al., 2017), a large-scale reading comprehension dataset with trivia-style factual questions.")

add_table(
    ["Property", "Value"],
    [
        ["Dataset", "TriviaQA"],
        ["Configuration", "rc (Reading Comprehension)"],
        ["Split", "Validation"],
        ["Questions processed", "1,000"],
        ["Library", "HuggingFace datasets"],
    ],
    "Table 7.4 — TriviaQA dataset configuration."
)

add_para("The rc configuration provides well-curated reference answers with comprehensive alias lists, essential for robust correctness evaluation. TriviaQA was selected because its short factual questions produce unambiguous ground truth — a prerequisite for evaluating hallucination detection accuracy.")

add_para("The TriviaQA rc configuration provides numerous fields per record: question, question_id, question_source, entity_pages (Wikipedia context with doc_source, filename, title, wiki_context), search_results (web search context with description, filename, rank, title, url, search_context), and answer (with sub-fields: value, aliases, normalized_value, normalized_aliases, matched_wiki_entity_name, normalized_matched_wiki_entity_name, type).")

add_para("Of these, only three fields are utilised:")

add_table(
    ["Field", "Type", "Purpose"],
    [
        ["question", "str", "Prompt input to the LLM"],
        ["question_id", "str", "Unique record identifier"],
        ["answer.aliases", "list[str]", "Reference answer variants for normalised substring correctness matching"],
    ],
    "Table 7.5 — TriviaQA fields used in dataset generation."
)

add_para("Crucially, no context documents are provided to the LLM — it answers purely from its parametric knowledge. This is intentional: the goal is to measure the model's internal confidence about its own knowledge, not its ability to extract answers from given passages.")

add_heading("7.3.3.2 Per-Question Record Generation", level=3)

add_para("For each of the 1,000 TriviaQA questions, the following pipeline produces a single training record:")

add_bullet("Multi-sample generation: ", bold_prefix="Step 1 — ")
add_para("Generate 5 independent responses using stochastic sampling (temperature = 0.7, max new tokens = 512), recording per-token logits and softmax probabilities for each sample.", indent=True)

add_bullet("Semantic clustering: ", bold_prefix="Step 2 — ")
add_para("Compare all sample pairs via LLM-based equivalence checking (greedy agglomerative clustering), producing cluster assignments.", indent=True)

add_bullet("Energy teacher computation: ", bold_prefix="Step 3 — ")
add_para("Apply Boltzmann energy aggregation (Formulas 2–3) and normalise across clusters to obtain energy_score_raw ∈ [0, 1].", indent=True)

add_bullet("Entropy teacher computation: ", bold_prefix="Step 4 — ")
add_para("Compute Shannon entropy (Formula 1) over cluster size distribution to obtain entropy_score_raw ∈ [0, ln 5].", indent=True)

add_bullet("Correctness evaluation: ", bold_prefix="Step 5 — ")
add_para("Check the main answer against all TriviaQA aliases using normalised substring matching, producing a binary correctness label. This label is used only for evaluation, not for training.", indent=True)

add_bullet("Hidden state extraction: ", bold_prefix="Step 6 — ")
add_para("Run a separate forward pass on the concatenation of prompt and main answer with output_hidden_states=True, extracting hidden states at the TBG position (last prompt token, shape 33 × 4096) and the SLT position (second-to-last generated token, shape 33 × 4096).", indent=True)

add_bullet("Logit feature extraction: ", bold_prefix="Step 7 — ")
add_para("Compute 9 scalar summary statistics from the main answer's per-token logits (mean, min, and standard deviation of chosen-token logits; mean, min, and standard deviation of top-1/top-2 margins; answer length; mean and min softmax probability).", indent=True)

add_para("Note: Hidden states are captured via a separate forward pass rather than during generation because storing hidden states during autoregressive generation would require 512 steps × 33 layers × 4096 dimensions ≈ 1.3 GB per sample — infeasible within GPU memory constraints.", italic=True)

add_figure_placeholder("Figure 7.2 — Flowchart: Per-question record generation pipeline showing the seven processing steps, which model calls occur at each step, and the data flowing between steps.")

add_para("The complete record schema is presented below.")

add_table(
    ["Field", "Type / Shape", "Description"],
    [
        ["uid", "str", "TriviaQA question ID"],
        ["question", "str", "Question text"],
        ["main_answer", "str", "First generated answer (sample 0)"],
        ["energy_score_raw", "float [0, 1]", "Energy teacher score (confidence)"],
        ["entropy_score_raw", "float [0, ~1.61]", "Entropy teacher score (uncertainty)"],
        ["correctness", "float {0.0, 1.0}", "Normalised substring match (evaluation only)"],
        ["emb_last_tok_before_gen", "ndarray (33, 4096)", "TBG hidden states — all 33 layers"],
        ["emb_tok_before_eos", "ndarray (33, 4096)", "SLT hidden states — all 33 layers"],
        ["logit_feats", "dict (9 keys)", "Token-level logit summary statistics"],
        ["num_clusters", "int", "Number of semantic clusters"],
        ["cluster_sizes", "list[int]", "Size of each cluster"],
    ],
    "Table 7.6 — Complete record schema in the generated probe training dataset."
)

add_para("After filtering records where the generated answer was too short for SLT extraction (fewer than 2 tokens), the final dataset comprises 500 valid records, saved as a serialised Python pickle file (probe_dataset_llama3-8b_triviaqa.pkl, approximately 541 MB).")

add_table(
    ["Split", "Records", "Percentage", "Purpose"],
    [
        ["Training", "400", "80%", "Binarisation thresholds, probe training"],
        ["Validation", "50", "10%", "Layer sweep AUROC, range selection"],
        ["Test", "50", "10%", "Final evaluation (held out)"],
    ],
    "Table 7.7 — Dataset split allocation (seed = 42 for reproducibility)."
)

# --- 7.3.4 ---
add_heading("7.3.4 Probe Training", level=2)

add_heading("7.3.4.1 Label Binarisation", level=3)

add_para("The raw teacher scores (energy and entropy) are continuous values. Since the probes are logistic regression classifiers, binary labels are required. The binarisation threshold is selected by sweeping from the 10th to the 90th percentile of the training distribution and choosing the value that minimises within-group mean squared error — the same approach used in Semantic Entropy Probes (Kossen et al., 2024).")

add_formula("y_i = 1  if  s_i ≥ τ,  else  0", "Formula 4")
add_para("where s_i is the raw teacher score and τ is the selected threshold. Thresholds are computed on the training split only and applied uniformly to validation and test splits.", italic=True)

add_table(
    ["Teacher Signal", "Threshold (τ)", "Label = 1 Interpretation"],
    [
        ["Energy", "0.7504", "High confidence (energy ≥ threshold)"],
        ["Entropy", "0.2052", "High uncertainty (entropy ≥ threshold)"],
    ],
    "Table 7.8 — Binarisation thresholds for Llama 3.1 8B on TriviaQA."
)

add_heading("7.3.4.2 Per-Layer AUROC Sweep", level=3)

add_para("To identify which transformer layers contain the strongest predictive signal, a per-layer sweep is conducted. For each of the 33 layers (32 transformer layers + 1 embedding layer), a logistic regression classifier (max iterations = 1000, regularisation C = 1.0) is trained on the single-layer hidden state vector (4,096 features) with StandardScaler preprocessing. Validation AUROC is recorded for each layer.")

add_para("This sweep is performed independently for all four probe types, revealing that different probes peak at different layers.")

add_table(
    ["Probe", "Best Single Layer", "Validation AUROC"],
    [
        ["TBG Energy", "Layer 14", "0.8706"],
        ["TBG Entropy", "Layer 24", "0.8693"],
        ["SLT Entropy", "Layer 20", "0.7929"],
        ["SLT Energy", "Layer 19", "0.7412"],
    ],
    "Table 7.9 — Best single-layer AUROC results from the per-layer sweep (Llama 3.1 8B)."
)

add_figure_placeholder("Figure 7.3 — Line chart: Per-layer AUROC sweep curves for all four probes across layers 0–32. TBG Energy peaks sharply at layer 14 (middle network), while TBG Entropy peaks at layer 24 (deeper). SLT probes show broader peaks in the upper-middle layers (19–20).")

add_heading("7.3.4.3 Layer Range Selection and Hidden State Semantics", level=3)

add_para("Rather than relying on a single peak layer, a contiguous window of layers is selected to improve robustness. Window sizes of 4, 8, and 16 layers are evaluated; the window with the highest mean validation AUROC is chosen.")

add_table(
    ["Probe", "Layer Range", "Window Size", "Mean Validation AUROC"],
    [
        ["Energy TBG", "(28, 32)", "4", "0.7840"],
        ["Entropy TBG", "(21, 25)", "4", "0.8523"],
        ["Entropy SLT", "(20, 24)", "4", "0.7725"],
        ["Energy SLT", "(17, 21)", "4", "0.7270"],
    ],
    "Table 7.10 — Selected layer ranges for the final probes (Llama 3.1 8B)."
)

add_para("The divergence in optimal layers reflects what different transformer layers encode:")

add_bullet("Early layers (0–8) ", bold_prefix="")
add_para("encode surface-level patterns: token identity, syntax, and basic word relationships. These layers recognise that a question is about geography but do not yet encode deep factual knowledge.", indent=True)

add_bullet("Middle layers (9–20) ", bold_prefix="")
add_para('encode semantic meaning and factual associations. This is where the model "retrieves" relevant knowledge from its training data. TBG Energy peaks here (layer 14) because factual confidence is most active in middle-layer representations.', indent=True)

add_bullet("Late layers (21–32) ", bold_prefix="")
add_para("encode output preparation: the model is resolving ambiguities, finalising token selection, and preparing the output distribution. TBG Entropy peaks here (layer 24) because uncertainty manifests in the output-preparation layers.", indent=True)

add_para("SLT probes operate on post-generation hidden states, which contain additional context from the generated answer. Their optimal ranges (layers 17–24) sit in the upper-middle region, where the model's accumulated generation state intersects with its confidence representations.")

add_heading("7.3.4.4 Final Probe Training and Evaluation", level=3)

add_para("The final probes are trained using logistic regression on the selected layer ranges. The feature vector for each record is formed by concatenating hidden state vectors across the selected layers and flattening: hidden[l₀:l₁, :].reshape(1, -1), producing a vector of (window_size × 4,096) features. StandardScaler normalisation is applied before training.")

add_para("Evaluation on the held-out test set with 1,000-sample bootstrap 95% confidence intervals:")

add_table(
    ["Probe", "Test AUROC", "95% CI"],
    [
        ["SLT Entropy", "0.7877", "[0.6547, 0.8949]"],
        ["TBG Entropy", "0.7857", "[0.6389, 0.8988]"],
        ["TBG Energy", "0.7480", "[0.5151, 0.9377]"],
        ["SLT Energy", "0.6667", "[0.4742, 0.8370]"],
    ],
    "Table 7.11 — Final probe test AUROC with bootstrap confidence intervals (Llama 3.1 8B)."
)

add_para("Feature ablation was conducted to evaluate whether logit summary features (mean/min/std of chosen-token logit, answer length) improve performance when combined with hidden states. The results confirm that hidden states alone are sufficient — adding logit features does not improve test AUROC. This is notable because logit features alone achieve high validation AUROC (~0.92), but this does not transfer to the test set, suggesting overfitting to the validation distribution.")

# --- 7.3.5 ---
add_heading("7.3.5 Teacher Fidelity and Cross-Signal Analysis", level=2)

add_heading("7.3.5.1 Teacher Fidelity", level=3)

add_para("Teacher fidelity measures how well each probe's continuous output reproduces the ranking of its corresponding teacher signal, assessed via Spearman rank correlation on the test set.")

add_table(
    ["Probe", "Spearman ρ", "p-value"],
    [
        ["SLT Entropy", "0.4330", "1.68 × 10⁻³"],
        ["TBG Entropy", "0.4329", "1.69 × 10⁻³"],
        ["TBG Energy", "0.2616", "6.64 × 10⁻²"],
        ["SLT Energy", "0.2507", "7.91 × 10⁻²"],
    ],
    "Table 7.12 — Teacher fidelity (Spearman ρ) between probe outputs and raw teacher scores."
)

add_para("Entropy probes exhibit higher fidelity (ρ ≈ 0.43) than energy probes (ρ ≈ 0.25). This is expected: entropy depends only on cluster sizes (a simpler structural signal), while energy additionally depends on per-token logit magnitudes within each cluster — a more complex pattern that a linear probe captures less faithfully.")

add_para("Crucially, low fidelity does not imply poor hallucination detection. The probes are trained on binarised labels (confident vs. uncertain), not on reproducing exact continuous rankings. A probe can correctly classify most questions as confident or uncertain (high AUROC) while imperfectly ordering them within each class (modest ρ).")

add_heading("7.3.5.2 Hallucination Detection Performance", level=3)

add_para("The practical measure of probe quality is how well they detect actual hallucinations, evaluated using model correctness as the ground truth.")

add_table(
    ["System", "Hallucination AUROC", "Cost"],
    [
        ["Full Energy Teacher (upper bound)", "0.7103", "60–120s"],
        ["Full Entropy Teacher (upper bound)", "0.7143", "60–120s"],
        ["SLT Energy Probe", "0.7163", "5–15s"],
        ["TBG Entropy Probe", "0.6806", "0.5–2s"],
        ["SLT Entropy Probe", "0.6726", "5–15s"],
        ["TBG Energy Probe", "0.6409", "0.5–2s"],
        ["Logit Features Only", "0.8075", "<1ms"],
    ],
    "Table 7.13 — Hallucination detection AUROC on the test set (Llama 3.1 8B, TriviaQA)."
)

add_para("A notable finding is that the SLT Energy probe (AUROC 0.7163) slightly exceeds the energy teacher upper bound (0.7103). This may be attributed to a regularisation benefit from binarisation — the probe learns a cleaner decision boundary by ignoring fine-grained score differences that are noise-dominated.")

add_heading("7.3.5.3 Cross-Signal Correlation", level=3)

add_para("At the teacher level, energy and entropy are near-perfectly anti-correlated:")

add_formula("ρ(energy_raw, entropy_raw) = −0.9958    (p = 1.48 × 10⁻⁵¹)", "")

add_para("This is expected: both are computed from the same semantic clustering — they are different mathematical summaries of the same underlying phenomenon. The remaining ~0.4% of unexplained variance corresponds to the logit-magnitude information that energy captures and entropy does not.")

add_para("At the probe level, the anti-correlation is substantially weaker:")

add_formula("ρ(energy_probe, entropy_probe) = −0.7444", "")

add_para("This divergence is significant: it demonstrates that the two probes extract partially independent information from hidden states, which justifies combining both signals in the per-sentence scoring pipeline (Section 7.3.7.2).")

# --- 7.3.6 ---
add_heading("7.3.6 Sentence-Level Decomposition and Claim Filtering", level=2)

add_para("LLM outputs typically contain multiple sentences, not all of which carry factual claims. Effective hallucination detection requires decomposing the output into individual sentences and identifying which ones warrant scoring.")

add_heading("7.3.6.1 Sentence Segmentation", level=3)

add_para('Sentence boundary detection is performed using pysbd (Pragmatic Sentence Boundary Disambiguation), a rules-based segmentation library. It was selected for three properties:')

add_bullet("Zero overhead: ", bold_prefix="")
add_para("No model loading, no GPU usage, deterministic output.", indent=True)
add_bullet("Robustness: ", bold_prefix="")
add_para('Handles abbreviations ("Dr.", "U.S."), decimal numbers ("2.1 million"), ellipsis, and numbered lists without false splits.', indent=True)
add_bullet("Configuration: ", bold_prefix="")
add_para("Initialised with language='en' and clean=False (preserves original whitespace and formatting).", indent=True)

add_heading("7.3.6.2 Token-to-Sentence Alignment", level=3)

add_para("After sentence segmentation, each generated token must be mapped to its parent sentence to enable per-sentence metric computation. The alignment procedure operates as follows:")

add_bullet("Tokenise the answer with return_offsets_mapping=True to obtain character spans for each token.")
add_bullet("Build sentence character spans by sequentially locating each sentence's start and end positions within the full text.")
add_bullet("For each token, compute the midpoint of its character span and assign it to the sentence whose character span contains that midpoint.")

add_para("This character-midpoint approach gracefully handles tokens that span sentence boundaries (e.g., punctuation tokens).")

add_heading("7.3.6.3 Claim Filtering", level=3)

add_para("The ClaimFilter class uses a pipeline of compiled regex patterns to classify each sentence as a factual claim or non-claim. Non-claim sentences (filler, hedging, meta-commentary) receive confidence = None and are excluded from all downstream aggregation.")

add_table(
    ["Category", "Pattern Description", "Example Match"],
    [
        ["Filler phrases", "Exact match against 11 common phrases", '"Of course", "Sure", "Thank you"'],
        ["Transitional/meta", 'Starts with "Here are...", "Let me...", "In summary..."', '"Here are the key points:"'],
        ["Headings/labels", "Short text (≤60 chars) ending with colon", '"Overview:", "Step 1:"'],
        ["Questions", "Starts with question word, ends with ?", '"What do you think?"'],
        ["Hedging/opinion", 'Starts with "I think...", "Perhaps...", "Maybe..."', '"I believe this is correct"'],
        ["Advisory/disclaimer", 'Starts with "Please note...", "Keep in mind..."', '"Please note that..."'],
        ["Greeting/sign-off", 'Starts with "Hi", "Hello", "Hope this helps..."', '"Hope this helps!"'],
        ["Offers to help", 'Starts with "If you need...", "Feel free..."', '"Feel free to ask"'],
        ["Enumeration intros", 'Starts with "The following...", "Some examples..."', '"Some key factors:"'],
        ["Markdown headers", "Bold text pattern (**text**:)", '"**Overview:**"'],
        ["Short non-numeric", "Fewer than 3 words with no digits", '"Okay."'],
    ],
    "Table 7.14 — ClaimFilter non-claim pattern categories with examples."
)

add_para("The filter operates conservatively: any sentence that does not match a non-claim pattern is treated as a factual claim and proceeds to scoring. This design prioritises recall (scoring all potential claims) over precision (potentially scoring some non-factual sentences).")

# --- 7.3.7 ---
add_heading("7.3.7 Per-Sentence Scoring Pipeline", level=2)

add_para("Each claim sentence is scored through two complementary mechanisms: a logit-based baseline and a probe-based risk estimate.")

add_heading("7.3.7.1 B1 Logit Confidence (Sigmoid Mapping)", level=3)

add_para("For each claim sentence, the mean chosen-token logit across all tokens in the sentence is computed and mapped to a [0, 1] confidence value via a calibrated sigmoid function.")

add_formula("confidence(s) = 1 / (1 + exp(−(mean_logit(s) − C) / S))", "Formula 5")
add_para("where C = 33.0 (sigmoid centre, calibrated for Llama 3.1 8B 4-bit quantised) and S = 3.0 (scale factor controlling curve steepness).", italic=True)

add_para("The centre value of 33.0 was calibrated to the typical logit range of the quantised model (approximately 25–45 for chosen tokens), placing the 50% confidence point at the distribution midpoint.")

add_para("Margin-based confidence boost: When the mean gap between the top-1 and top-2 logits exceeds 3.0, a linear boost is applied:", bold=True)

add_formula("boost = min(0.1,  (mean_margin − 3.0) × 0.02)", "")
add_formula("confidence(s) = min(1.0,  confidence(s) + boost)", "")

add_para("This rewards sentences where the model exhibited strong commitment to its chosen tokens.")

add_heading("7.3.7.2 Per-Sentence Dual-Probe Scoring", level=3)

add_para("For each claim sentence in SLT mode, the hidden state at the last token position of that sentence is extracted and passed through both the SLT energy and SLT entropy probes.")

add_bullet("Energy probe: Outputs P(confident). Risk is inverted: energy_risk = 1 − predict_proba(X)[0, 1].")
add_bullet("Entropy probe: Outputs P(uncertain). Risk is used directly: entropy_risk = predict_proba(X)[0, 1].")

add_para("The two risks are combined using weights derived from their relative validation AUROC performance:")

add_formula("probe_risk(s) = W_entropy · entropy_risk(s) + W_energy · energy_risk(s)", "Formula 6")

add_para("where W_entropy = 0.773 / (0.773 + 0.727) = 0.515 and W_energy = 0.727 / (0.773 + 0.727) = 0.485.", italic=True)

add_para("The entropy probe receives slightly higher weight (0.515 vs 0.485) because it demonstrated higher discrimination on the validation set (AUROC 0.773 vs 0.727).")

add_para("Risk-to-level override: The probe risk can escalate (but never reduce) the B1 confidence level:", bold=True)
add_bullet("If probe_risk ≥ 0.65 → override level to \"low\" (high hallucination risk)")
add_bullet("If probe_risk ≥ 0.35 → override level to \"medium\"")
add_bullet("If probe_risk < 0.35 → retain the original B1 level")

# --- 7.3.8 ---
add_heading("7.3.8 Aggregate Hallucination Score", level=2)

add_para("The final step combines per-sentence and overall probe scores into a single aggregate risk score. The aggregation strategy adapts to answer length, reflecting the insight that the SLT hidden state becomes less representative as the generated sequence grows longer.")

add_heading("7.3.8.1 Token-Length Conditional Aggregation", level=3)

add_para('A threshold of 100 tokens separates "short" and "long" answers. This boundary was calibrated against the TriviaQA training distribution, where most correct factual answers fall below this length.')

add_para("Short Answer Aggregation (≤ 100 tokens):", bold=True)

add_formula("slt_combined = (energy_risk + entropy_risk) / 2", "Formula 7a")
add_formula("combined_risk = 0.5 · slt_combined + 0.5 · mean(R_sent)    if n_claims ≥ 2", "Formula 7b")
add_formula("combined_risk = slt_combined    if n_claims < 2", "Formula 7c")

add_para("For short answers with 0–1 claims, the SLT probe alone is sufficient since the SLT token position captures the entire answer's context. For 2 or more claims, a 50-50 blend with the per-sentence average provides coverage across all claims.")

add_para("Long Answer Aggregation (> 100 tokens):", bold=True)

add_formula("w_slt = 0.15", "")
add_formula("w_max = 0.25 / (1 + ln(max(n, 1)))", "Formula 9")
add_formula("w_mean = 1.0 − w_slt − w_max", "")
add_formula("combined_risk = w_slt · entropy_risk + w_max · max(R_sent) + w_mean · mean(R_sent)", "Formula 8")

add_para("The design rationale for each weight component:")
add_bullet("SLT weight (0.15): ", bold_prefix="")
add_para("For long answers, the SLT token position only reflects the end of the generated sequence, making it less representative. Only the entropy component is used. It serves as a lightweight anchor rather than a primary signal.", indent=True)
add_bullet("Max-sentence weight (Formula 9): ", bold_prefix="")
add_para("Decays logarithmically with the number of claims, preventing a single high-risk sentence from dominating the score when many claims are present.", indent=True)
add_bullet("Mean-sentence weight: ", bold_prefix="")
add_para("The remainder, always the largest component, providing a stable centre-of-mass risk across all claims.", indent=True)

add_table(
    ["Number of Claims (n)", "w_max", "w_mean"],
    [
        ["1", "0.250", "0.600"],
        ["5", "0.096", "0.754"],
        ["10", "0.075", "0.775"],
        ["20", "0.063", "0.787"],
    ],
    "Table 7.15 — Weight distribution for long-answer aggregation as claim count increases."
)

add_heading("7.3.8.2 Final Confidence Level Mapping", level=3)

add_table(
    ["Combined Risk", "Confidence Level", "Interpretation"],
    [
        ["< 0.35", "High", "Low hallucination risk"],
        ["< 0.65", "Medium", "Moderate hallucination risk"],
        ["≥ 0.65", "Low", "High hallucination risk"],
    ],
)

add_table(
    ["Answer Type", "SLT Weight", "Per-Sentence Role", "Max-Sentence Role"],
    [
        ["Short (≤ 100 tokens), 0–1 claims", "1.0", "Not used", "Not used"],
        ["Short (≤ 100 tokens), 2+ claims", "0.5", "0.5 (mean)", "Not used"],
        ["Long (> 100 tokens)", "0.15 (entropy only)", "Remainder (mean)", "0.25 / ln(n+1)"],
    ],
    "Table 7.16 — Summary of aggregation strategies by answer type."
)

# --- 7.3.9 ---
add_heading("7.3.9 System Architecture and Module Design", level=2)

add_para("This section covers the system's structure, code design, and module integration. Table 7.16 first delineates which components are adopted from prior work and which are original contributions of this project.")

add_table(
    ["Component", "Origin", "Notes"],
    [
        ["Semantic Entropy framework", "Kuhn et al. (2023)", "Multi-sample generation, semantic clustering, Shannon entropy formula"],
        ["Semantic Energy teacher", "Prior work", "Boltzmann energy per sample, cluster energy normalisation"],
        ["pysbd sentence segmentation", "Open-source library", "Adopted unchanged; language='en', clean=False"],
        ["sklearn LogisticRegression", "Open-source library", "Used as probe classifier; hyperparameters tuned in this work"],
        ["Probe architecture (4-probe matrix) [THIS WORK]", "This work", "TBG/SLT position rationale, 2x2 teacher x position design, layer range selection"],
        ["Hidden state extraction pipeline [THIS WORK]", "This work", "Separate forward pass design, tensor shape handling, sentence-end token positions"],
        ["Label binarisation scheme [THIS WORK]", "This work", "MSE-minimising threshold sweep on training distribution"],
        ["ClaimFilter [THIS WORK]", "This work", "Sequential regex pipeline, 9 non-claim pattern categories"],
        ["Conditional risk aggregation [THIS WORK]", "This work", "Token-length conditional strategy, logarithmic max-sentence decay (Formula 8)"],
        ["FastAPI scoring API [THIS WORK]", "This work", "Three-endpoint architecture (Full SE / Fast SLT / Fast TBG)"],
        ["Frontend SPA [THIS WORK]", "This work", "Vanilla JS chat interface with per-sentence risk bars and SVG score history"],
    ],
    "Table 7.16 — Contribution scope: adopted components and original contributions (marked [THIS WORK])."
)

add_heading("7.3.9.1 Backend Architecture", level=3)

add_para("The backend follows a modular design with clear separation of concerns.")

add_table(
    ["Module", "File", "Responsibility"],
    [
        ["API Layer", "backend/app.py (297 lines)", "FastAPI application, endpoint routing, model lifecycle management, probe bundle loading"],
        ["Semantic Engine", "backend/engine.py (671 lines)", "LLM loading (8-bit quantised), generation, semantic clustering, hidden state extraction, probe scoring, sentence-level scoring, risk aggregation"],
        ["Claim Filter", "backend/claim_filter.py (74 lines)", "Regex-based sentence classification (claim vs non-claim), zero external dependencies"],
        ["Probe Bundle", "backend/models/*.pkl", "Serialised bundle: 4 LogisticRegression probes, 4 StandardScaler preprocessors, 4 layer range tuples"],
    ],
    "Table 7.17 — Backend module responsibilities."
)

add_para("The SemanticEngine class is the core computational module, responsible for model management, response generation, semantic analysis, hidden state extraction, probe inference, and sentence processing.")

add_heading("7.3.9.2 SemanticEngine Code Structure", level=3)

add_para("The SemanticEngine class in backend/engine.py is instantiated once at application startup and held as a module-level variable in app.py — a Singleton pattern that keeps the LLM (~8 GB VRAM) resident across requests, eliminating the 30-60 second reload cost on every query.")

add_table(
    ["Method", "Lines", "Purpose"],
    [
        ["__init__", "51-93", "CUDA validation, BitsAndBytes config, model/tokeniser loading, pysbd and ClaimFilter init, probe bundle loading"],
        ["generate_responses", "~94-160", "Stochastic multi-sample generation; returns per-token logits and softmax probabilities"],
        ["check_equivalence", "~161-220", "LLM-based pairwise semantic equivalence; greedy decoding"],
        ["cluster_responses", "~221-270", "Greedy agglomerative clustering using pairwise equivalence matrix"],
        ["score_sentences", "~271-350", "pysbd segmentation, claim filtering, B1 logit confidence per sentence"],
        ["_extract_hidden_states", "351-401", "Separate forward pass; TBG/SLT hidden state extraction; sentence-end token positions"],
        ["score_with_tbg_probe", "403-459", "TBG probe inference: layer slice -> StandardScaler -> predict_proba -> risk mapping"],
        ["score_with_slt_probe", "461-540+", "SLT probe inference: generation + hidden extraction + per-sentence scoring + aggregation"],
        ["run_full_pipeline", "~541-671", "Full Semantic Energy: 5 generations, clustering, energy/entropy computation"],
    ],
    "Table 7.18 — SemanticEngine method map with approximate line references (backend/engine.py)."
)

add_para("Singleton initialisation — CUDA check and 8-bit model loading (engine.py lines 51-72):")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "class SemanticEngine:\n"
    "    def __init__(self, model_name, probe_bundle_path):\n"
    "        if not torch.cuda.is_available():\n"
    "            raise RuntimeError('SemanticEngine requires CUDA.')\n"
    "        bnb_config = BitsAndBytesConfig(load_in_8bit=True, llm_int8_threshold=6.0)\n"
    "        self.model = AutoModelForCausalLM.from_pretrained(\n"
    "            model_name, quantization_config=bnb_config, device_map='auto')\n"
    "        self.tokenizer    = AutoTokenizer.from_pretrained(model_name)\n"
    "        self.segmenter    = pysbd.Segmenter(language='en', clean=False)\n"
    "        self.claim_filter = ClaimFilter()\n"
    "        with open(probe_bundle_path, 'rb') as f:\n"
    "            self.probe_bundle = pickle.load(f)"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_para("The BitsAndBytes 8-bit configuration reduces Llama 3.1 8B from ~16 GB (FP16) to ~8 GB VRAM. The probe bundle is co-loaded with the model to enforce structural consistency across model switches.")

add_para("Core energy computation — Boltzmann teacher (engine.py lines 39-48):")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "def cal_boltzmann_logits(logits_list):\n"
    "    return [-mean(sublist) for sublist in logits_list]\n\n"
    "def cal_flow(probs_list, logits_list, clusters, fermi_mu=None):\n"
    "    probs = cal_probs(probs_list)\n"
    "    logits = (cal_fermi_dirac_logits(logits_list, mu=fermi_mu)\n"
    "              if fermi_mu is not None else cal_boltzmann_logits(logits_list))\n"
    "    return cal_cluster_ce(probs, logits, clusters)"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_para("cal_boltzmann_logits implements Formula 2 (negative mean chosen-token logit = per-sample energy). cal_flow is the central teacher computation: it assembles probabilities, selects the energy variant, and delegates cluster aggregation and normalisation to cal_cluster_ce (Formula 3).")

add_para("Hidden state extraction — separate forward pass (engine.py lines 378-401):")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "with torch.no_grad():\n"
    "    outputs = self.model(**full_inputs, output_hidden_states=True)\n\n"
    "hidden = torch.stack(outputs.hidden_states, dim=0)  # (layers, 1, seq_len, hidden_dim)\n"
    "hidden = hidden[:, 0, :, :].float().cpu()           # (layers, seq_len, hidden_dim)\n\n"
    "tbg_hidden = hidden[:, prompt_len - 1, :].numpy()  # last prompt token\n"
    "slt_hidden = hidden[:, full_len - 2, :].numpy()    # second-to-last generated token"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_para("torch.stack converts the per-layer tuple into a single tensor. Index full_len-2 targets the second-to-last token (EOS is always last). Both arrays move to CPU before probe inference to free VRAM.")

add_para("Probe inference — layer slice and logistic regression (engine.py lines 431-451):")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "l0, l1 = probe_bundle['best_energy_tbg_range']\n"
    "X_e   = probe_bundle['tbg_energy_scaler'].transform(tbg_hidden[l0:l1].reshape(1,-1))\n"
    "energy_risk = 1.0 - probe_bundle['tbg_energy_probe'].predict_proba(X_e)[0, 1]\n\n"
    "l0h, l1h = probe_bundle['best_entropy_tbg_range']\n"
    "X_h   = probe_bundle['tbg_entropy_scaler'].transform(tbg_hidden[l0h:l1h].reshape(1,-1))\n"
    "entropy_risk = probe_bundle['tbg_entropy_probe'].predict_proba(X_h)[0, 1]\n\n"
    "combined_risk = (energy_risk + entropy_risk) / 2.0"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_para("Layer range tuples in the probe bundle ensure inference always uses the exact layers trained. reshape(1,-1) flattens (window, 4096) into the vector expected by predict_proba. Energy probe outputs P(confident) and is inverted to risk; entropy probe outputs P(uncertain) directly.")

add_heading("7.3.9.3 API Endpoints", level=3)

add_para("Three scoring endpoints implement the three speed-accuracy tradeoff points:")

add_table(
    ["Endpoint", "Method", "Scoring Mode", "Model Calls", "Response Time"],
    [
        ["/chat", "POST", "Full Semantic Energy", "5 generations + up to 10 equivalence checks", "60-120s"],
        ["/score_fast_slt", "POST", "Fast SLT (probe-based)", "1 generation + 1 forward pass", "5-15s"],
        ["/score_fast_tbg", "POST", "Fast TBG (probe-based)", "1 forward pass (no generation)", "0.5-2s"],
        ["/status", "GET", "Health check", "None", "< 1ms"],
        ["/switch_model", "POST", "Model swap", "Full model reload", "30-60s"],
    ],
    "Table 7.19 — API endpoints and their characteristics."
)

add_figure_placeholder("Figure 7.4 — System architecture diagram showing three tiers: Frontend (Vanilla JS), Backend (FastAPI with SemanticEngine, ClaimFilter, and ProbeBundle), and ML Layer (LLM with Hidden State Extractor). Arrows indicate data flow between components.")

add_figure_placeholder("Figure 7.5 — Sequence diagram: Side-by-side comparison of the three scoring modes. Full SE performs all steps; Fast SLT skips clustering; Fast TBG skips both generation and clustering.")

add_heading("7.3.9.4 Training Pipeline", level=3)

add_para("The training infrastructure consists of four Jupyter notebooks forming a sequential pipeline:")

add_table(
    ["Notebook", "Purpose", "Output"],
    [
        ["00_preflight.ipynb", "Formula verification, score orientation checks", "Validated mathematical foundations"],
        ["01_generate_dataset.ipynb", "TriviaQA processing, 5-sample generation, teacher computation", "probe_dataset_llama3-8b_triviaqa.pkl (~541 MB)"],
        ["02_train_se_probes.ipynb", "Binarisation, layer sweep, probe training, evaluation", "probes_llama3-8b_triviaqa.pkl (~2.1 MB)"],
        ["04_sentence_baseline.ipynb", "B1 per-sentence logit confidence validation", "Calibrated sigmoid parameters"],
    ],
    "Table 7.20 — Training pipeline notebooks and their outputs."
)

add_heading("7.3.9.5 Frontend Architecture", level=3)

add_para("The frontend is implemented as a vanilla JavaScript single-page application with no framework dependencies. Key components include:")

add_bullet("Three-mode selector (Full SE / Fast SLT / Fast TBG) persisted in localStorage, allowing users to select their preferred speed-accuracy tradeoff.")
add_bullet("Per-message rendering with a colour-coded confidence badge (green/yellow/red), response timer, and collapsible metrics panel showing per-sentence analysis with risk bars.")
add_bullet("Score history chart (SVG-based) visualising confidence trends across the conversation, with colour-coded zones for high/medium/low confidence.")
add_bullet("Model switching via dropdown, with a loading overlay during model swap.")

add_para("State management: localStorage stores persistent preferences (mode, model selection); sessionStorage stores ephemeral state (chat message HTML, score history).")

add_heading("7.3.9.6 Module Integration", level=3)

add_para("Method-level call sequence for /score_fast_slt (the endpoint with the most integration touch-points):")

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "POST /score_fast_slt  (app.py)\n"
    "  |\n"
    "  +-- Readiness check: engine not None, probe_bundle loaded\n"
    "  |\n"
    "  +-- engine.score_with_slt_probe(prompt, max_new_tokens=512)\n"
    "       |\n"
    "       +-- 1. generate_responses(n=1, temperature=0.7)\n"
    "       |       -> answer_text, logits_per_token, probs_per_token\n"
    "       |\n"
    "       +-- 2. score_sentences(answer_text, logits_per_token)\n"
    "       |       +-- pysbd.Segmenter.segment()     -> sentence spans\n"
    "       |       +-- ClaimFilter.is_claim()         -> bool per sentence\n"
    "       |       +-- B1 sigmoid(mean_logit, margin) -> confidence per claim\n"
    "       |\n"
    "       +-- 3. _align_tokens_to_sentences()        -> token->sentence mapping\n"
    "       |\n"
    "       +-- 4. _extract_hidden_states(prompt_ids, answer_ids)\n"
    "       |       +-- model(**full_inputs, output_hidden_states=True)\n"
    "       |       +-- tbg_hidden = [:, prompt_len-1, :]  (33, 4096)\n"
    "       |       +-- slt_hidden = [:, full_len-2, :]    (33, 4096)\n"
    "       |       +-- sentence_end_positions\n"
    "       |\n"
    "       +-- 5. SLT probe inference (overall)\n"
    "       |       energy_risk  = 1 - slt_energy_probe.predict_proba()[:,1]\n"
    "       |       entropy_risk = slt_entropy_probe.predict_proba()[:,1]\n"
    "       |\n"
    "       +-- 6. Per-sentence probe scoring\n"
    "       |       probe_risk(s) = 0.515*entropy_risk + 0.485*energy_risk\n"
    "       |       level(s)      = max(B1_level, probe_risk_level)\n"
    "       |\n"
    "       +-- 7. Conditional aggregation (Formula 7 or 8)\n"
    "               -> combined_risk, confidence_level\n"
    "               -> JSON {answer, confidence_level, combined_risk, ...}"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_para("Integration design decisions:")
add_bullet("Singleton engine: app.py holds the engine as a module-level variable initialised at startup. Every request handler accesses the same instance — the correct pattern for a single-GPU server where model loading is the dominant cost.")
add_bullet("Atomic model-probe binding: When /switch_model is called, the engine is deleted, CUDA cache cleared, and a new engine instantiated with the model-specific probe bundle path. This enforces probe-model consistency structurally: running a Qwen3 probe against a Llama hidden state is impossible because the bundle is always co-loaded.")
add_bullet("Claim filter decoupling: ClaimFilter holds only compiled re patterns and has no GPU dependencies. It operates purely on strings, so sentence classification never blocks on GPU operations.")
add_bullet("TBG as a pre-generation gate: /score_fast_tbg runs a forward pass on the prompt alone, returning a risk estimate before the model generates any answer. This enables a qualitatively different interaction mode: the frontend can warn the user before potentially hallucinated text is shown.")

# --- 7.3.10 ---
add_heading("7.3.10 Functional Requirement Traceability", level=2)

add_para("The following table maps each functional requirement to its implementing module(s).")

add_table(
    ["ID", "Requirement", "Implementation", "Module(s)"],
    [
        ["FR1", "Analyse Answer (Local)", "All inference on local GPU via 8-bit quantised LLM; no external API calls", "engine.py (SemanticEngine.__init__)"],
        ["FR2", "Compute Paragraph Risk", "combined_risk and confidence_level in all three scoring modes", "engine.py, app.py"],
        ["FR3", "Factoid/Sentence Decomposition", "pysbd segmentation + character-midpoint token alignment", "engine.py (score_sentences, align_tokens_to_sentences)"],
        ["FR4", "Per-Factoid Risk Labels", "Per-sentence probe_risk, level, and is_claim in SLT and Full SE modes", "engine.py (score_sentences, score_with_slt_probe)"],
        ["FR5", "Single-Pass Probe Estimator", "TBG: 1 forward pass; SLT: 1 generation + 1 forward pass", "engine.py (score_with_tbg_probe, score_with_slt_probe)"],
        ["FR6", "Risk Summary View", "Colour-coded confidence badge + overall score per message", "frontend/script.js"],
        ["FR7", "Detailed Rationale", "Collapsible metrics panel with per-sentence breakdown", "frontend/script.js"],
        ["FR8", "Threshold & Display Settings", "Mode selector (Full SE / SLT / TBG) with speed-accuracy tradeoffs", "frontend/script.js, app.py"],
        ["FR9", "Multi-user Accounts / Auth", "Not implemented — out of scope", "—"],
        ["FR10", "Public Web/App Hosting", "Not implemented — local deployment only", "—"],
        ["FR11", "Multilingual Detection", "Not implemented — English only", "—"],
    ],
    "Table 7.21 — Functional requirement traceability matrix."
)

add_para("All mandatory requirements (FR1-FR5, priority M) are fully implemented. The 'should have' requirement (FR6) and 'could have' requirement (FR7) are also implemented. FR8 is partially addressed through the three-mode selection mechanism. FR9-FR11 are explicitly documented as out of scope.")

# -- Save --
output_path = os.path.join(os.path.dirname(__file__), "thesis_7_3_core_functionalities.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
